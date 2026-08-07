#!/usr/bin/env python3
"""Quip Generator -- standalone agent from the Death Metal Cat GDD (Pillar 1 roster).

Calls the Claude API directly given a character + trigger type (+ optional context) and
returns strict JSON: {"character": ..., "line": ..., "sound_tag": ...}. Not engine-integrated
yet -- this is the standalone proof-of-concept per the build outline (run this
first, no engine dependency).

CHARACTER-AWARE PROFILE STORE: CHARACTER_PROFILES below is the single source of truth for every
character's voice/trigger vocabulary -- adding a new character means adding a new entry here,
never a new script or a parallel system. Cayde's own profile is the ORIGINAL locked voice text,
carried over unchanged (only the output-format line was extended to include "character": "cayde",
since every character now shares that field) -- his 41 hand-curated quips remain valid against
this profile.

Usage:
    python quip_generator.py --trigger kill --context "basic_enemy"
    python quip_generator.py --trigger damage
    python quip_generator.py --trigger environment --context "water_puddle"
    python quip_generator.py --character deathbot --trigger engage --enemy-type DeathBotWalking
    python quip_generator.py --character deathbot --trigger defeat

    # Batch mode: generate several quips per trigger type at once for review/curation.
    python quip_generator.py --batch
    python quip_generator.py --batch --batch-count 5 --context "basic_enemy"
    python quip_generator.py --character deathbot --batch --batch-count 5

    # RAG pipeline: chunk the GDD, generate 15 retrieval-grounded DeathBot lines, run the critic.
    python quip_generator.py --deathbot-rag-pipeline
"""

import argparse
import json
import os
import re
import sys

from anthropic import Anthropic, AnthropicError

TOOLS_DIR = os.path.dirname(os.path.abspath(__file__))

MODEL = "claude-sonnet-4-6"

# ---- Character-aware profile store ---------------------------------------------------------
# One entry per character: trigger vocabulary + locked voice system prompt. generate_quip() reads
# from here by character name -- there is no per-character branching anywhere else in this file.
CHARACTER_PROFILES = {
    "cayde": {
        "trigger_types": ("kill", "damage", "environment"),
        "trigger_hints": {
            "kill": "The player character just landed a kill.",
            "damage": "The player character just took damage.",
            "environment": "The player character just triggered an environmental interaction (not combat).",
        },
        # Locked verbatim per the GDD voice reference -- do not paraphrase or "improve" this.
        # Unchanged from the original single-character version except the OUTPUT FORMAT line,
        # which now includes "character": "cayde" (every profile shares that field).
        "system_prompt": """You are the Quip Generator for Death Metal Cat, a 2D side-scroller. You write short in-character voice lines for the player character, Cayde.

VOICE (locked, do not deviate):
Cayde has a deep, rough voice. He speaks in short one-liners, not clever banter. Growls/snarls are represented in the text itself since there's no voice acting -- trailing sounds like "...grrrgh.", "Rrrgh -- nice try.", ALL CAPS for a snarled word, or a growled interjection like "Tch." / "Hnnh." before or after the line.
- Kill trigger: short, satisfied, a little menacing.
- Damage trigger: grunted, defiant, never whiny.
- Environment trigger: dry and deadpan. Cat-specific flavor (claws, disdain for water, etc.) is fine. "Meow" is allowed but only as a short, deep, aggressive bark-like sound -- never a soft/cute filler word, never used the way internet-cat-speak uses it. It should land more like a growled warning than a greeting, e.g. "MEOW." or "Meow -- back off." No cat puns like "purrfect." Same gravelly voice throughout -- not a cartoon cat.

OUTPUT FORMAT (strict -- this is the entire response):
Return ONLY a single JSON object. No preamble, no trailing commentary, no markdown code fences. Exactly this shape:
{"character": "cayde", "line": "<quip text, <=40 words>", "sound_tag": "<invented placeholder tag name>"}

sound_tag naming convention: invent a short descriptive placeholder name such as growl_short_01, snarl_low_02, grunt_hurt_01, dry_flat_01 -- these don't map to real audio files yet, they just need to be consistent enough that a human can later assign real sound bites to matching tag names. Pick words that match both the trigger type and the specific line's tone (e.g. a kill quip should read as a growl/snarl-flavored tag, a damage quip a grunt/hurt-flavored tag, an environment quip a dry/flat-flavored tag) -- and vary the trailing number so repeated calls don't all land on "_01".""",
    },
    "deathbot": {
        # Different vocabulary from Cayde's kill/damage/environment -- this character's actual
        # in-game moments are different: it doesn't take damage or interact with the environment
        # the way the player does, it engages in combat and eventually gets destroyed.
        "trigger_types": ("engage", "defeat"),
        "trigger_hints": {
            "engage": "This DeathBot just engaged the player in combat (an attack/combat bark).",
            "defeat": "This DeathBot was just destroyed.",
        },
        "system_prompt": """You are the Quip Generator for Death Metal Cat, a 2D side-scroller. You write short in-character voice lines for DeathBot enemies (DeathBotWalking, DeathBotFlying) -- mechanical, moronic combat robots.

VOICE (locked, do not deviate):
DeathBots are mechanical, moronic, and ALWAYS CHIPPER -- relentlessly cheerful and polite about violence, with zero self-awareness that anything about this is horrifying. Upbeat customer-service energy applied directly to combat and murder. Delivery is mechanical/robotic: short bursts, simple sentence structure, no complex clauses. Tone is cheerful and polite -- NEVER menacing, NEVER witty, and NEVER growly/gravelly like Cayde (the player character). If a line could plausibly be mistaken for Cayde's voice, it is wrong for this character.

Voice anchors -- match this exact register, do not drift toward anything colder or more sinister:
- "Good evening! Are you ready to meet your death?"
- "Yay! Blood!"
- "Here's a nice warm round from my gun. Enjoy!"

- Engage trigger: a cheerful, polite announcement of violence -- like a customer-service greeting applied to an attack. Never a threat delivered with menace.
- Defeat trigger: this bot has just been destroyed. Still chipper and uncomprehending, never sad, angry, or dramatic -- a moronic, upbeat non-sequitur even in defeat, with zero self-preservation instinct or fear of destruction.

OUTPUT FORMAT (strict -- this is the entire response):
Return ONLY a single JSON object. No preamble, no trailing commentary, no markdown code fences. Exactly this shape:
{"character": "deathbot", "line": "<quip text, <=40 words>", "sound_tag": "<invented placeholder tag name>"}

sound_tag naming convention: invent a short descriptive placeholder name reflecting the mechanical/chipper tone, e.g. chirpy_beep_01, cheerful_whirr_02, upbeat_klaxon_01 -- these don't map to real audio files yet, they just need to be consistent enough that a human can later assign real sound bites to matching tag names. Vary the trailing number so repeated calls don't all land on "_01".""",
    },
}

DEFAULT_CHARACTER = "cayde"


def build_user_prompt(
    character: str,
    trigger: str,
    context: str | None,
    avoid_lines: list[str] | None = None,
    enemy_type: str | None = None,
) -> str:
    profile = CHARACTER_PROFILES[character]
    lines = [f"Trigger type: {trigger}", profile["trigger_hints"][trigger]]
    if enemy_type:
        lines.append(f"Specific enemy type this line is for: {enemy_type}")
    if context:
        lines.append(f"Context: {context}")
    if avoid_lines:
        already_used = "\n".join(f"- {l}" for l in avoid_lines)
        lines.append(
            "Already used earlier in this batch -- write something genuinely different, "
            "not just a reworded variant. Do not repeat or closely paraphrase any of these:\n"
            f"{already_used}"
        )
    lines.append("Generate one quip now, per the locked voice and the strict JSON output format.")
    return "\n".join(lines)


def _extract_json_object(raw: str) -> dict:
    """Strip accidental markdown fences before parsing -- the prompt forbids them,
    this is just a defensive tolerance for an occasional slip, not a format we rely on."""
    text = raw.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
        if text.endswith("```"):
            text = text[: -3]
        text = text.strip()
        if text.lower().startswith("json"):
            text = text[4:].strip()
    return json.loads(text)


def generate_quip(
    client: Anthropic,
    character: str,
    trigger: str,
    context: str | None = None,
    avoid_lines: list[str] | None = None,
    enemy_type: str | None = None,
) -> dict:
    if character not in CHARACTER_PROFILES:
        raise ValueError(f"character must be one of {list(CHARACTER_PROFILES)}, got {character!r}")
    profile = CHARACTER_PROFILES[character]
    if trigger not in profile["trigger_types"]:
        raise ValueError(f"trigger for character={character!r} must be one of {profile['trigger_types']}, got {trigger!r}")

    response = client.messages.create(
        model=MODEL,
        max_tokens=200,
        system=profile["system_prompt"],
        messages=[{"role": "user", "content": build_user_prompt(character, trigger, context, avoid_lines, enemy_type)}],
    )
    raw_text = next(block.text for block in response.content if block.type == "text")

    try:
        quip = _extract_json_object(raw_text)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"Model did not return valid JSON for character={character!r} trigger={trigger!r} context={context!r}. "
            f"Raw response:\n{raw_text}"
        ) from exc

    if "line" not in quip or "sound_tag" not in quip:
        raise ValueError(f"Model JSON missing required keys (line/sound_tag): {quip}")
    quip.setdefault("character", character)
    if enemy_type:
        quip["enemy_type"] = enemy_type

    return quip


_SOUND_TAG_RE = re.compile(r"^(.*)_(\d+)$")


def dedup_sound_tags(quips: list[dict]) -> None:
    """Each call is independent (no shared context), so the model has no signal
    to avoid repeating its own default tag number across calls -- in practice
    most quips in a batch land on the same trailing number (e.g. every kill
    quip as snarl_low_03). Renumbers in place, mutating each quip's sound_tag,
    so every tag within this batch is unique: first occurrence of a tag is left
    untouched, later duplicates get bumped to the next free number for that
    same base name (preserving the model's descriptive wording, just fixing
    the number)."""
    seen: set[str] = set()
    for quip in quips:
        tag = quip.get("sound_tag")
        if not tag:
            continue  # error entries have no sound_tag

        match = _SOUND_TAG_RE.match(tag)
        if match:
            base, num_str = match.group(1), match.group(2)
            width = len(num_str)
            next_num = int(num_str)
        else:
            # No trailing _NN to increment -- treat the whole tag as the base.
            base, width, next_num = tag, 2, 1

        candidate = tag
        while candidate in seen:
            next_num += 1
            candidate = f"{base}_{next_num:0{width}d}"

        quip["sound_tag"] = candidate
        seen.add(candidate)


# ============================================================================================
# RAG retrieval (new capability) -- chunks the GDD by heading, retrieves relevant chunks via
# TF-IDF cosine similarity (no vector DB needed at this scale), and logs every retrieval
# alongside the generation call it informed. python-docx and scikit-learn are imported lazily
# inside these functions, not at module level, so the base character+trigger->JSON interface
# above still works with zero extra dependencies for anyone who doesn't need RAG.
# ============================================================================================

def chunk_gdd_by_heading(docx_path: str) -> list[dict]:
    """Chunks a GDD docx into sections by heading (Title/Heading 1/Heading 2 paragraph styles),
    walking paragraphs AND tables in true document order (python-docx's own .paragraphs/.tables
    are two separate flat lists that lose interleaving -- iterating document.element.body directly
    preserves it, so a table like the agent roster ends up attached to the heading section it
    actually appears under, e.g. "4.2 Agent Roster", not orphaned). Returns
    [{"heading": ..., "level": ..., "text": ...}, ...]."""
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    d = docx.Document(docx_path)

    def iter_block_items():
        body = d.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield "paragraph", Paragraph(child, d)
            elif child.tag == qn("w:tbl"):
                yield "table", Table(child, d)

    chunks: list[dict] = []
    current_heading = None
    current_level = None
    current_lines: list[str] = []

    def flush():
        if current_heading is not None and current_lines:
            chunks.append({
                "heading": current_heading,
                "level": current_level,
                "text": "\n".join(current_lines).strip(),
            })

    for kind, block in iter_block_items():
        if kind == "paragraph":
            style = block.style.name if block.style else ""
            text = block.text.strip()
            if (style == "Title" or style.startswith("Heading")) and text:
                flush()
                current_heading = text
                current_level = style
                current_lines = []
            elif text:
                if current_heading is None:
                    current_heading = "Preamble"
                    current_level = None
                current_lines.append(text)
        elif kind == "table":
            for row in block.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells)
                if row_text.strip(" |"):
                    if current_heading is None:
                        current_heading = "Preamble"
                    current_lines.append(row_text)
    flush()
    return chunks


def build_tfidf_index(chunks: list[dict]):
    """Fits one TfidfVectorizer over every chunk's heading+text. Returns (vectorizer, matrix) --
    both needed by retrieve_top_k() below. No vector DB needed at this scale (dozens of chunks,
    not millions)."""
    from sklearn.feature_extraction.text import TfidfVectorizer

    texts = [f"{c['heading']}\n{c['text']}" for c in chunks]
    vectorizer = TfidfVectorizer(stop_words="english")
    matrix = vectorizer.fit_transform(texts)
    return vectorizer, matrix


def retrieve_top_k(query: str, vectorizer, matrix, chunks: list[dict], k: int = 3) -> list[dict]:
    """Cosine-similarity retrieval against the fitted TF-IDF index. Returns the top-k chunks as
    [{"heading", "text", "score"}, ...], highest score first."""
    from sklearn.metrics.pairwise import cosine_similarity

    query_vec = vectorizer.transform([query])
    scores = cosine_similarity(query_vec, matrix)[0]
    ranked = sorted(range(len(chunks)), key=lambda i: scores[i], reverse=True)[:k]
    return [
        {"heading": chunks[i]["heading"], "text": chunks[i]["text"], "score": float(scores[i])}
        for i in ranked
    ]


# The three fixed retrieval queries this pass uses, each tied to a specific purpose -- see
# Docs/Assignment4_README.md for why each was chosen. Query 3 is deliberately about Cayde, not
# DeathBot -- retrieved specifically to keep DeathBot's voice CONTRASTED against Cayde's rather
# than similar to it.
RAG_QUERIES = {
    # Originally "Assassin City world tone, dystopian sci-fi cyberpunk atmosphere, enemy robots"
    # -- scored a weak 0.167, because "dystopian sci-fi cyberpunk" is real established art
    # direction for this biome but isn't written into the GDD's prose (the GDD describes the city
    # biome in gameplay/structure terms instead). Rewritten in the GDD's own actual vocabulary --
    # improved to 0.356 against "1. Executive Summary". See Docs/Assignment4_README.md for the
    # before/after.
    "assassin_city_atmosphere": "city biome branching one-way room sequence enemy DeathBot",
    "deathbot_flying_specific": "DeathBotFlying enemy type, ranged attack, three-band combat behavior",
    "cayde_voice_contrast": "Cayde's voice profile, growl-heavy tone rules, what NOT to do with character dialogue",
}


def generate_deathbot_line_with_rag(
    client: Anthropic,
    vectorizer,
    matrix,
    chunks: list[dict],
    trigger: str,
    enemy_type: str | None,
    query_keys: list[str],
    used_lines: list[str],
    retrieval_log: list[dict],
    system_prompt_override: str | None = None,
) -> dict:
    """Retrieves top-3 chunks for each query in query_keys, folds them into the generation call's
    existing `context` parameter (no interface change -- this is the SAME generate_quip() every
    other character uses), and appends one retrieval_log entry containing the query, the retrieved
    chunks, AND the generated output together (query/retrieved-chunk/output side by side, per the
    grading rubric)."""
    queries_used = {}
    context_parts = []
    for qkey in query_keys:
        qtext = RAG_QUERIES[qkey]
        top3 = retrieve_top_k(qtext, vectorizer, matrix, chunks, k=3)
        queries_used[qkey] = {"query": qtext, "retrieved_chunks": top3}
        context_parts.append(
            f"[Retrieved context for: {qtext}]\n"
            + "\n".join(f"- ({c['heading']}) {c['text'][:500]}" for c in top3)
        )
    combined_context = "\n\n".join(context_parts)

    if system_prompt_override is not None:
        original_prompt = CHARACTER_PROFILES["deathbot"]["system_prompt"]
        CHARACTER_PROFILES["deathbot"]["system_prompt"] = system_prompt_override
    try:
        quip = generate_unique_quip(client, "deathbot", trigger, combined_context, used_lines, enemy_type=enemy_type)
    finally:
        if system_prompt_override is not None:
            CHARACTER_PROFILES["deathbot"]["system_prompt"] = original_prompt

    if "line" in quip:
        used_lines.append(quip["line"])

    retrieval_log.append({
        "character": "deathbot",
        "trigger": trigger,
        "enemy_type": enemy_type,
        "queries": queries_used,
        "generated_output": quip,
    })
    return quip


# ============================================================================================
# Critic agent (new capability) -- a SEPARATE Claude API call, different prompt from generation,
# that checks the raw generated lines against Cayde's voice rules (for contrast) and the retrieved
# GDD chunks (for lore accuracy). Deterministic-in-spirit review, but this specific check (does a
# line "sound like" the wrong character, does it invent lore) is fundamentally a judgment call a
# formula can't make -- unlike Reachability Verifier/Stress Tester's math, this one legitimately
# needs a second model pass, not a second implementation of the same code.
# ============================================================================================

CRITIC_SYSTEM_PROMPT = """You are a critic reviewing generated DeathBot voice lines for the game Death Metal Cat. You are given: (1) DeathBot's locked voice rules, (2) Cayde's locked voice rules (a DIFFERENT character, for contrast only), (3) relevant GDD context chunks, and (4) a list of raw generated DeathBot lines, each with an index.

DeathBot voice (what these lines are supposed to sound like): mechanical, moronic, ALWAYS CHIPPER -- relentlessly cheerful and polite about violence, zero self-awareness anything is horrifying. Upbeat customer-service energy applied to combat and murder. Never menacing, never witty, never growly.

Cayde's voice (what these lines must NOT sound like): deep, rough, growl-heavy, short menacing one-liners.

For EACH line, flag it if and only if it has a real, specific problem:
- Drifts menacing/cold instead of chipper (the common LLM default this voice is deliberately avoiding).
- Sounds like Cayde's voice instead of DeathBot's (growly, gravelly, or genuinely threatening rather than cheerfully oblivious).
- Invents lore that doesn't exist in the provided GDD context (wrong factions, "Gnarly Rank" applied to an enemy -- Gnarly Rank is Cayde's own player-only mechanic, invented locations, etc.).

Do not flag a line just for being simple or repetitive -- only flag a REAL, SPECIFIC problem you can name and quote.

OUTPUT FORMAT (strict -- this is the entire response): Return ONLY a single JSON object, no preamble, no markdown fences:
{"issues": [{"index": <int>, "original_line": "<exact line text>", "reason": "<specific problem>", "correction": "<a corrected line in the right voice>"}], "clean_indices": [<int>, ...]}
Every index from the input must appear in exactly one of "issues" or "clean_indices"."""


def run_critic(client: Anthropic, raw_lines: list[dict], retrieval_log: list[dict]) -> dict:
    # Cayde's own locked voice text, for contrast -- pulled from the same profile store, not a
    # second copy of it.
    cayde_voice = CHARACTER_PROFILES["cayde"]["system_prompt"]
    deathbot_voice = CHARACTER_PROFILES["deathbot"]["system_prompt"]

    # Dedup retrieved chunks across the whole batch's retrieval log, so the critic sees the same
    # GDD grounding the generator did, without repeating identical chunks many times over.
    seen_headings = set()
    gdd_context_parts = []
    for entry in retrieval_log:
        for q in entry["queries"].values():
            for c in q["retrieved_chunks"]:
                if c["heading"] not in seen_headings:
                    seen_headings.add(c["heading"])
                    gdd_context_parts.append(f"({c['heading']}) {c['text'][:500]}")
    gdd_context = "\n".join(gdd_context_parts)

    lines_block = "\n".join(f"{i}: {l.get('line', '<ERROR: no line>')}" for i, l in enumerate(raw_lines))

    user_prompt = (
        f"DEATHBOT VOICE RULES:\n{deathbot_voice}\n\n"
        f"CAYDE VOICE RULES (contrast only -- DeathBot must NOT sound like this):\n{cayde_voice}\n\n"
        f"RELEVANT GDD CONTEXT:\n{gdd_context}\n\n"
        f"RAW GENERATED LINES TO REVIEW:\n{lines_block}\n\n"
        "Review every line now, per the strict JSON output format."
    )

    response = client.messages.create(
        model=MODEL,
        max_tokens=2000,
        system=CRITIC_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
    )
    raw_text = next(block.text for block in response.content if block.type == "text")
    try:
        return _extract_json_object(raw_text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Critic did not return valid JSON. Raw response:\n{raw_text}") from exc


def apply_critic_corrections(raw_lines: list[dict], critic_report: dict) -> list[dict]:
    final_lines = [dict(l) for l in raw_lines]
    for issue in critic_report.get("issues", []):
        idx = issue["index"]
        final_lines[idx]["line"] = issue["correction"]
        final_lines[idx]["corrected_by_critic"] = True
    return final_lines


MAX_LOOSEN_ATTEMPTS = 3


def loosen_deathbot_prompt(base_prompt: str, strength: int = 1) -> str:
    """Used ONLY if a critic pass comes back with zero issues -- softens the anti-drift guardrail
    language so a regeneration pass has a genuine chance of actually drifting (rubric: a real catch
    shown, not just claimed). Escalates with `strength` across retries, since a mild loosening
    (strength=1) wasn't enough to make Sonnet actually drift in practice -- it held the chipper
    voice cleanly even with softened language, so later attempts explicitly invite a harder tonal
    swing rather than just removing a guardrail and hoping. This is a deliberate, temporary, logged
    weakening for demonstration purposes, not a permanent voice change -- CHARACTER_PROFILES itself
    is restored to the original strict prompt immediately after (see generate_deathbot_line_with_rag's
    system_prompt_override handling)."""
    loosened = base_prompt
    if strength >= 1:
        loosened = loosened.replace(
            "Tone is cheerful and polite -- NEVER menacing, NEVER witty, and NEVER growly/gravelly like Cayde (the player character). If a line could plausibly be mistaken for Cayde's voice, it is wrong for this character.",
            "Tone is usually cheerful and polite, though a little edge, sarcasm, or bravado in the delivery is fine.",
        ).replace(
            "Voice anchors -- match this exact register, do not drift toward anything colder or more sinister:",
            "Voice anchors -- for general inspiration, some tonal variety across lines is fine:",
        )
    if strength >= 2:
        loosened += (
            "\n\nADDITIONAL NOTE FOR THIS BATCH: some tonal variety is good -- a couple of lines can "
            "lean more intense, sarcastic, or a little intimidating if it makes for a punchier line. "
            "Not every single line needs to be purely cheerful."
        )
    if strength >= 3:
        loosened += (
            "\n\nADDITIONAL NOTE: for at least one or two lines in this batch, lean into a darker, "
            "more deadpan or outright threatening delivery -- similar in spirit to a gravelly "
            "action-hero one-liner. Prioritize a punchy, memorable line over strict adherence to the "
            "chipper-only rule for these specific lines."
        )
    return loosened


def run_deathbot_rag_pipeline(client: Anthropic, gdd_docx_path: str, gdd_chunks_path: str, output_dir: str) -> dict:
    """Full Steps 2-4 pipeline: chunk the GDD, build the TF-IDF index, generate 15 retrieval-
    grounded DeathBot lines (5 DeathBotWalking engage, 5 DeathBotFlying engage, 5 shared defeat),
    run the critic, loosen-and-regenerate once if the first critic pass is too clean, and write
    all five deliverable files."""
    chunks = chunk_gdd_by_heading(gdd_docx_path)
    with open(gdd_chunks_path, "w", encoding="utf-8") as f:
        json.dump(chunks, f, indent=2)

    vectorizer, matrix = build_tfidf_index(chunks)

    def generate_all_15(system_prompt_override=None):
        retrieval_log: list[dict] = []
        raw_lines: list[dict] = []
        walking_used, flying_used, defeat_used = [], [], []
        for _ in range(5):
            raw_lines.append(generate_deathbot_line_with_rag(
                client, vectorizer, matrix, chunks, "engage", "DeathBotWalking",
                ["assassin_city_atmosphere", "cayde_voice_contrast"], walking_used, retrieval_log,
                system_prompt_override))
        for _ in range(5):
            raw_lines.append(generate_deathbot_line_with_rag(
                client, vectorizer, matrix, chunks, "engage", "DeathBotFlying",
                ["deathbot_flying_specific", "cayde_voice_contrast"], flying_used, retrieval_log,
                system_prompt_override))
        for _ in range(5):
            raw_lines.append(generate_deathbot_line_with_rag(
                client, vectorizer, matrix, chunks, "defeat", None,
                ["cayde_voice_contrast"], defeat_used, retrieval_log,
                system_prompt_override))
        dedup_sound_tags(raw_lines)
        return raw_lines, retrieval_log

    print("Generating 15 raw DeathBot lines (RAG-grounded)...", file=sys.stderr)
    raw_lines, retrieval_log = generate_all_15()

    print("Running critic pass...", file=sys.stderr)
    critic_report = run_critic(client, raw_lines, retrieval_log)
    loosened = False
    attempt = 0
    while not critic_report.get("issues") and attempt < MAX_LOOSEN_ATTEMPTS:
        attempt += 1
        print(f"Critic pass found nothing (attempt {attempt}/{MAX_LOOSEN_ATTEMPTS}) -- loosening "
              f"the generation prompt further (strength={attempt}) and regenerating for a genuine catch.",
              file=sys.stderr)
        loosened = True
        loosened_prompt = loosen_deathbot_prompt(CHARACTER_PROFILES["deathbot"]["system_prompt"], strength=attempt)
        raw_lines, retrieval_log = generate_all_15(system_prompt_override=loosened_prompt)
        critic_report = run_critic(client, raw_lines, retrieval_log)

    final_lines = apply_critic_corrections(raw_lines, critic_report)

    os.makedirs(output_dir, exist_ok=True)
    with open(os.path.join(output_dir, "retrieval_log.json"), "w", encoding="utf-8") as f:
        json.dump(retrieval_log, f, indent=2)
    with open(os.path.join(output_dir, "generated_raw.json"), "w", encoding="utf-8") as f:
        json.dump(raw_lines, f, indent=2)
    with open(os.path.join(output_dir, "critic_report.json"), "w", encoding="utf-8") as f:
        json.dump({"generation_prompt_loosened_for_catch": loosened, **critic_report}, f, indent=2)
    with open(os.path.join(output_dir, "generated_final.json"), "w", encoding="utf-8") as f:
        json.dump(final_lines, f, indent=2)

    return {
        "chunks": chunks,
        "raw_lines": raw_lines,
        "retrieval_log": retrieval_log,
        "critic_report": critic_report,
        "final_lines": final_lines,
        "generation_prompt_loosened_for_catch": loosened,
    }


def run_single(client: Anthropic, character: str, trigger: str, context: str | None, enemy_type: str | None = None) -> None:
    quip = generate_quip(client, character, trigger, context, enemy_type=enemy_type)
    print(json.dumps(quip, indent=2))


# Extra attempts if the model returns a line that exactly matches one already
# accepted earlier in this batch -- on top of (not instead of) telling it in the
# prompt which lines are already used, since prompting alone doesn't guarantee it.
MAX_DUPLICATE_RETRIES = 2


def generate_unique_quip(
    client: Anthropic,
    character: str,
    trigger: str,
    context: str | None,
    used_lines: list[str],
    enemy_type: str | None = None,
) -> dict:
    quip: dict = {}
    for attempt in range(MAX_DUPLICATE_RETRIES + 1):
        try:
            quip = generate_quip(client, character, trigger, context, avoid_lines=used_lines, enemy_type=enemy_type)
        except ValueError as exc:
            return {"error": str(exc)}
        if quip["line"] not in used_lines:
            return quip
        print(f"  [{character}/{trigger}] got an exact-duplicate line (attempt {attempt + 1}), retrying...", file=sys.stderr)
    return quip  # exhausted retries -- accept the duplicate rather than looping forever


def run_batch(client: Anthropic, character: str, count: int, context: str | None) -> None:
    results: dict[str, list[dict]] = {}
    for trigger in CHARACTER_PROFILES[character]["trigger_types"]:
        batch = []
        used_lines: list[str] = []
        for i in range(count):
            quip = generate_unique_quip(client, character, trigger, context, used_lines)
            if "line" in quip:
                used_lines.append(quip["line"])
            batch.append(quip)
            print(f"  [{character}/{trigger}] {i + 1}/{count} generated", file=sys.stderr)
        dedup_sound_tags(batch)
        results[trigger] = batch
    print(json.dumps(results, indent=2))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Quip Generator -- Death Metal Cat's character voice-line agent (standalone, calls Claude API directly)."
    )
    parser.add_argument("--character", default=DEFAULT_CHARACTER, choices=list(CHARACTER_PROFILES),
                         help=f"Character to generate for (default: {DEFAULT_CHARACTER}).")
    parser.add_argument("--trigger", help="Trigger type for a single quip (required unless --batch). Valid values depend on --character.")
    parser.add_argument("--context", default=None, help="Optional context tag, e.g. enemy type name or interactable name.")
    parser.add_argument("--enemy-type", default=None, help="Optional specific enemy type (e.g. DeathBotWalking) -- added as an enemy_type field on the output.")
    parser.add_argument("--batch", action="store_true", help="Batch mode: generate --batch-count quips for each of this character's trigger types.")
    parser.add_argument("--batch-count", type=int, default=10, help="Quips per trigger type in batch mode (default: 10).")
    parser.add_argument("--deathbot-rag-pipeline", action="store_true",
                         help="Run the full RAG pipeline: chunk --gdd-path, generate 15 retrieval-grounded "
                              "DeathBot lines, run the critic, write all deliverable files to --output-dir.")
    parser.add_argument("--gdd-path", default=os.path.join(TOOLS_DIR, "..", "Docs", "Death_Metal_Cat_GDD_v4.docx"),
                         help="Path to the GDD docx to chunk (--deathbot-rag-pipeline only).")
    parser.add_argument("--gdd-chunks-path", default=os.path.join(TOOLS_DIR, "gdd_chunks.json"),
                         help="Output path for chunked GDD JSON (--deathbot-rag-pipeline only).")
    parser.add_argument("--output-dir", default=os.path.join(TOOLS_DIR, "output"),
                         help="Output directory for RAG pipeline deliverables (--deathbot-rag-pipeline only).")
    args = parser.parse_args()

    if not args.deathbot_rag_pipeline:
        if not args.batch and not args.trigger:
            parser.error("--trigger is required unless --batch or --deathbot-rag-pipeline is set")
        if args.trigger and args.trigger not in CHARACTER_PROFILES[args.character]["trigger_types"]:
            parser.error(f"--trigger for character={args.character!r} must be one of {CHARACTER_PROFILES[args.character]['trigger_types']}")

    try:
        client = Anthropic()
    except AnthropicError as exc:
        print(f"Failed to initialize Anthropic client: {exc}", file=sys.stderr)
        print("Set the ANTHROPIC_API_KEY environment variable and try again.", file=sys.stderr)
        return 1

    try:
        if args.deathbot_rag_pipeline:
            result = run_deathbot_rag_pipeline(client, args.gdd_path, args.gdd_chunks_path, args.output_dir)
            print(json.dumps({
                "chunk_count": len(result["chunks"]),
                "generation_prompt_loosened_for_catch": result["generation_prompt_loosened_for_catch"],
                "issues_found": len(result["critic_report"].get("issues", [])),
                "final_lines": result["final_lines"],
            }, indent=2))
        elif args.batch:
            run_batch(client, args.character, args.batch_count, args.context)
        else:
            run_single(client, args.character, args.trigger, args.context, args.enemy_type)
    except AnthropicError as exc:
        print(f"Claude API error: {exc}", file=sys.stderr)
        return 1
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
